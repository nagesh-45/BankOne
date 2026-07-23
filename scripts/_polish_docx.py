#!/usr/bin/env python3
"""Polish all BankOne docs/*.docx for consistent margins and table alignment."""
from pathlib import Path
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH


def set_cell_margins(cell, top=40, bottom=40, left=60, right=60):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for m, v in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        node = OxmlElement(f"w:{m}")
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def style_table(table):
    table.autofit = True
    table.allow_autofit = True
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:type"), "pct")
    tblW.set(qn("w:w"), "5000")
    old = tblPr.find(qn("w:tblBorders"))
    if old is not None:
        tblPr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "CBD5E0")
        borders.append(el)
    tblPr.append(borders)
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_margins(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.1
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Calibri"
            if r_idx == 0:
                tcPr = cell._tc.get_or_add_tcPr()
                shd = tcPr.find(qn("w:shd"))
                if shd is None:
                    shd = OxmlElement("w:shd")
                    tcPr.append(shd)
                shd.set(qn("w:fill"), "1A365D")
                shd.set(qn("w:val"), "clear")
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.font.size = Pt(9)


def polish(path: Path) -> None:
    doc = Document(str(path))
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
    for p in doc.paragraphs:
        if p.style and p.style.name == "Normal":
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.15
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for table in doc.tables:
        style_table(table)
    doc.save(str(path))


def main() -> None:
    root = Path(sys.argv[1] if len(sys.argv) > 1 else ".")
    for path in sorted(root.glob("*.docx")):
        if path.name.startswith("_"):
            continue
        polish(path)
        print("polished", path.name)
    modules = root / "MODULES"
    if modules.is_dir():
        for path in sorted(modules.glob("*.docx")):
            polish(path)
            print("polished MODULES/" + path.name)


if __name__ == "__main__":
    main()
